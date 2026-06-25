"""
Run Full Workflow
=================
Executes all workflow steps defined in project_config.json, runs strict
release-gate checks, writes a versioned workflow execution audit, and
prints the final release decision.

Step types supported
--------------------
check_path      : verify a file or directory exists
command         : run an arbitrary shell command
qa_status       : read the Status: line from a QA markdown file
render_outputs  : built-in Markdown -> DOCX + PDF renderer
qa_script       : run a QA Python script from <skill_dir>/scripts/qa/
                  The script receives --project-root <project_root> plus any
                  extra args listed in the step's "args" list.

Release decisions
-----------------
MACHINE_QA_FAILED                       one or more machine gates failed
MACHINE_QA_PASSED_HUMAN_PENDING         machine gates passed; author/expert items remain
SUBMISSION_READY_AFTER_AUTHOR_CONFIRMATION   all gates clear
"""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

SKILL_DIR = Path(__file__).resolve().parents[1]
HUMAN_STATUSES = {"NEEDS_AUTHOR", "NEEDS_EXTERNAL_EXPERT", "NEEDS_HUMAN_REVIEW"}


# ── Helpers ───────────────────────────────────────────────────────────────────

def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def artifact_status(path: Path) -> str:
    """Read the first 'Status: <value>' line from a QA markdown file.

    Normalizes the verdict token to uppercase and extracts only the first word
    after 'Status:', so 'Status: pass -- see notes' and 'Status: PASS' both
    resolve to 'PASS'.  This makes the gate robust to human editing.
    """
    if not path.exists():
        return "MISSING"
    for line in path.read_text(encoding="utf-8", errors="ignore").splitlines()[:40]:
        stripped = line.strip()
        if stripped.lower().startswith("status:"):
            raw = stripped.split(":", 1)[1].strip()
            # Extract the first token only (handles trailing comments/dashes/notes)
            token = (raw.split()[0].upper()) if raw.split() else "PRESENT_STATUS_NOT_DECLARED"
            return token if token else "PRESENT_STATUS_NOT_DECLARED"
    return "PRESENT_STATUS_NOT_DECLARED"


def next_audit_version(output_dir: Path) -> int:
    """Return the next unused vN index for the workflow execution audit."""
    existing = [
        int(m.group(1))
        for f in output_dir.glob("workflow_execution_audit_v*.json")
        if (m := re.search(r"v(\d+)\.json$", f.name))
    ]
    return (max(existing) + 1) if existing else 1


def run_subprocess(project_root: Path, command: list[str], timeout: int = 180) -> dict:
    proc = subprocess.run(
        command,
        cwd=str(project_root),
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout,
    )
    return {
        "command": command,
        "returncode": proc.returncode,
        "stdout": proc.stdout[-4000:],
        "stderr": proc.stderr[-4000:],
        "status": "PASS" if proc.returncode == 0 else "FAIL",
    }


# ── DOCX / PDF renderer ───────────────────────────────────────────────────────

def _apply_inline(para_obj, text: str) -> None:
    """Apply bold and italic inline formatting to a python-docx paragraph."""
    # Pattern: **bold**, *italic*, ***bold-italic***
    tokens = re.split(r"(\*{1,3}[^*]+\*{1,3})", text)
    for token in tokens:
        if token.startswith("***") and token.endswith("***"):
            run = para_obj.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith("**") and token.endswith("**"):
            run = para_obj.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = para_obj.add_run(token[1:-1])
            run.italic = True
        else:
            para_obj.add_run(token)


def _parse_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into a list of row-lists."""
    rows = []
    for line in lines:
        if re.match(r"^\|[-: |]+\|$", line.strip()):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def render_basic_outputs(project_root: Path, config: dict) -> dict:
    outputs = config.get("outputs", {})
    inputs = config.get("inputs", {})
    manuscript_path = project_root / outputs.get(
        "manuscript_md", inputs.get("manuscript", "01_manuscript/manuscript.md")
    )
    docx_path = project_root / outputs.get("docx", "02_outputs/manuscript.docx")
    pdf_path = project_root / outputs.get("pdf", "02_outputs/manuscript.pdf")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    if not manuscript_path.exists():
        return {
            "name": "render_basic_outputs",
            "type": "render_outputs",
            "status": "FAIL",
            "error": f"Manuscript not found: {manuscript_path}",
        }

    text = manuscript_path.read_text(encoding="utf-8", errors="ignore")

    # ── DOCX ──────────────────────────────────────────────────────────────────
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt

        document = Document()

        # Default style
        normal = document.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

        # Line numbers for the default section
        try:
            sect_pr = document.sections[0]._sectPr
            ln_num = OxmlElement("w:lnNumType")
            ln_num.set(qn("w:countBy"), "1")
            ln_num.set(qn("w:restart"), "newPage")
            sect_pr.append(ln_num)
        except Exception:
            pass  # line numbers are optional; don't block the render

        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headings
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
                i += 1
                continue
            if line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
                i += 1
                continue
            if line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
                i += 1
                continue

            # Table: collect consecutive | lines
            if line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                rows = _parse_table(table_lines)
                if rows:
                    tbl = document.add_table(rows=len(rows), cols=len(rows[0]))
                    tbl.style = "Table Grid"
                    for r_idx, row_cells in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_cells):
                            if c_idx < len(tbl.rows[r_idx].cells):
                                tbl.rows[r_idx].cells[c_idx].text = cell_text
                continue

            # Unordered list item
            if re.match(r"^\s*[-*+]\s+", line):
                content = re.sub(r"^\s*[-*+]\s+", "", line)
                p = document.add_paragraph(style="List Bullet")
                _apply_inline(p, content)
                i += 1
                continue

            # Ordered list item
            if re.match(r"^\s*\d+\.\s+", line):
                content = re.sub(r"^\s*\d+\.\s+", "", line)
                p = document.add_paragraph(style="List Number")
                _apply_inline(p, content)
                i += 1
                continue

            # Blank line -> paragraph break (collect next non-blank run)
            if not line.strip():
                i += 1
                continue

            # Regular paragraph: accumulate until blank line or structural break
            block_lines = []
            while i < len(lines):
                l = lines[i]
                if not l.strip():
                    break
                if l.startswith("#") or l.strip().startswith("|") or re.match(r"^\s*[-*+\d]", l):
                    break
                block_lines.append(l.strip())
                i += 1
            if block_lines:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 2.0  # double-spacing for submission
                _apply_inline(p, " ".join(block_lines))

        document.save(docx_path)

    except Exception as exc:
        return {
            "name": "render_basic_outputs",
            "type": "render_outputs",
            "status": "FAIL",
            "error": f"DOCX render failed: {exc}",
        }

    # ── PDF ───────────────────────────────────────────────────────────────────
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.lib import colors

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "BodyDouble",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=12,
            leading=24,  # double-spaced
            spaceAfter=6,
        )
        h1_style = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Times-Bold", fontSize=14)
        h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Times-Bold", fontSize=12)

        story = []
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            if line.startswith("# "):
                story.append(Paragraph(line[2:].strip(), h1_style))
                story.append(Spacer(1, 6))
                i += 1
                continue
            if line.startswith("## "):
                story.append(Paragraph(line[3:].strip(), h2_style))
                story.append(Spacer(1, 4))
                i += 1
                continue
            if line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                rows = _parse_table(table_lines)
                if rows:
                    pdf_table = Table(rows)
                    pdf_table.setStyle(TableStyle([
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 8))
                continue
            if not line.strip():
                i += 1
                continue
            block_lines = []
            while i < len(lines):
                l = lines[i]
                if not l.strip() or l.startswith("#") or l.strip().startswith("|"):
                    break
                block_lines.append(l.strip())
                i += 1
            if block_lines:
                clean = " ".join(block_lines)
                # Escape XML special chars for ReportLab
                clean = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                # Convert **bold** and *italic* to ReportLab tags
                clean = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", clean)
                clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", clean)
                clean = re.sub(r"\*(.+?)\*", r"<i>\1</i>", clean)
                story.append(Paragraph(clean, body_style))

        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=LETTER,
            leftMargin=1.25 * inch,
            rightMargin=1.25 * inch,
            topMargin=1.0 * inch,
            bottomMargin=1.0 * inch,
        )
        doc.build(story)

    except Exception as exc:
        return {
            "name": "render_basic_outputs",
            "type": "render_outputs",
            "status": "FAIL",
            "error": f"PDF render failed: {exc}",
        }

    return {
        "name": "render_basic_outputs",
        "type": "render_outputs",
        "status": "PASS",
        "docx": str(docx_path),
        "pdf": str(pdf_path),
    }


# ── Step dispatcher ───────────────────────────────────────────────────────────

def run_step(project_root: Path, step: dict, config: dict) -> dict:
    kind = step.get("type", "check_path")
    name = step["name"]

    if kind == "render_outputs":
        result = render_basic_outputs(project_root, config)
        result["name"] = name
        return result

    if kind == "check_path":
        path = project_root / step["path"]
        return {"name": name, "type": kind, "path": str(path),
                "status": "PASS" if path.exists() else "FAIL"}

    if kind == "command":
        result = run_subprocess(project_root, step["command"], int(step.get("timeout", 180)))
        result["name"] = name
        result["type"] = kind
        return result

    if kind == "qa_script":
        # Resolve the script from the skill's scripts/qa/ directory
        script_name = step["script"]
        script_path = SKILL_DIR / "scripts" / "qa" / script_name
        if not script_path.exists():
            return {"name": name, "type": kind, "status": "FAIL",
                    "error": f"QA script not found: {script_path}"}
        cmd = [sys.executable, str(script_path), "--project-root", str(project_root)]
        cmd += step.get("args", [])
        result = run_subprocess(project_root, cmd, int(step.get("timeout", 120)))
        result["name"] = name
        result["type"] = kind
        return result

    if kind == "qa_status":
        path = project_root / step["path"]
        observed = artifact_status(path)
        required = step.get("required_status", "PASS")
        return {
            "name": name, "type": kind, "path": str(path),
            "required_status": required, "observed_status": observed,
            "status": "PASS" if observed == required else "FAIL",
        }

    return {"name": name, "type": kind, "status": "FAIL",
            "error": f"Unknown step type: {kind}"}


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description="Run the full Therasik academic writing workflow")
    parser.add_argument("project_config", help="Path to project_config.json")
    parser.add_argument("--skip-qa-scripts", action="store_true",
                        help="Skip qa_script steps (for framework testing only)")
    args = parser.parse_args()

    config_path = Path(args.project_config).resolve()
    project_root = config_path.parent
    config = read_json(config_path)
    started = datetime.now(timezone.utc).isoformat()

    skip_qa = args.skip_qa_scripts
    steps = [run_step(project_root, step, config) for step in config.get("workflow_steps", [])
             if not (skip_qa and step.get("type") == "qa_script")]
    failed_steps = [s for s in steps if s.get("status") not in {"PASS", "PASS_BASIC"}]

    machine_gate_failures: list[str] = []
    human_gate_pending: list[str] = []
    gate_items: list[dict] = []

    for gate_name, gate in config.get("strict_release_artifacts", {}).items():
        path = project_root / gate["path"]
        required = gate.get("required_status", "PASS")
        observed = artifact_status(path)
        human_boundary = bool(gate.get("human_boundary", False)) or observed in HUMAN_STATUSES
        passed = observed == required
        item = {
            "gate": gate_name, "path": str(path),
            "required_status": required, "observed_status": observed,
            "human_boundary": human_boundary, "pass": passed,
            "purpose": gate.get("purpose", ""),
        }
        gate_items.append(item)
        if not passed:
            msg = f"{gate_name}: observed={observed}; required={required}"
            (human_gate_pending if human_boundary else machine_gate_failures).append(msg)

    automated_status = "PASS" if not failed_steps and not machine_gate_failures else "FAIL"
    if automated_status == "FAIL":
        release_decision = "MACHINE_QA_FAILED"
    elif human_gate_pending:
        release_decision = "MACHINE_QA_PASSED_HUMAN_PENDING"
    else:
        release_decision = "SUBMISSION_READY_AFTER_AUTHOR_CONFIRMATION"

    output_dir = project_root / config.get("outputs", {}).get("qa_dir", "03_QA")
    output_dir.mkdir(parents=True, exist_ok=True)

    version = next_audit_version(output_dir)
    audit = {
        "project_id": config.get("project_id"),
        "generated_utc": datetime.now(timezone.utc).isoformat(),
        "started_utc": started,
        "audit_version": version,
        "config": str(config_path),
        "automated_status": automated_status,
        "release_decision": release_decision,
        "submission_ready": automated_status == "PASS" and not human_gate_pending,
        "steps": steps,
        "failed_steps": failed_steps,
        "strict_release_artifacts": gate_items,
        "machine_gate_failures": machine_gate_failures,
        "human_gate_pending": human_gate_pending,
    }

    # Write versioned audit file
    audit_path = output_dir / f"workflow_execution_audit_v{version}.json"
    audit_path.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")

    # Update latest pointer
    latest_ptr = output_dir / "workflow_status.json"
    latest_ptr.write_text(
        json.dumps({"latest_audit": audit_path.name, "release_decision": release_decision},
                   indent=2),
        encoding="utf-8",
    )

    print(json.dumps(audit, indent=2, ensure_ascii=False))
    return 0 if release_decision != "MACHINE_QA_FAILED" else 1


if __name__ == "__main__":
    raise SystemExit(main())
