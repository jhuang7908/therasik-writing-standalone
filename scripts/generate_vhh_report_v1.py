"""
VHH  v1.0

：
-  result.json 
-  Markdown 
-  .md  .docx 
- 

：
- python-docx ( .docx )
- markdown ( Markdown )
"""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. .docx export will be disabled.")

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False
    print("Warning: markdown not installed. Markdown processing will be limited.")


def load_template(template_path: Path) -> str:
    """ Markdown """
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    return template_path.read_text(encoding="utf-8")


def extract_sequence_analysis(result: Dict[str, Any]) -> Dict[str, Any]:
    """ result """
    seq_analysis = result.get("sequence_analysis", {})
    orig_regions = seq_analysis.get("original_regions", {}) or {}
    
    #  IMGT （， numbering ）
    regions_info = {}
    current_pos = 1
    
    for region_name in ["FR1", "CDR1", "FR2", "CDR2", "FR3", "CDR3", "FR4"]:
        seq = orig_regions.get(region_name, "")
        if seq:
            regions_info[f"{region_name.lower()}_seq"] = seq
            regions_info[f"{region_name.lower()}_len"] = len(seq)
            regions_info[f"{region_name.lower()}_start"] = current_pos
            current_pos += len(seq)
    
    return regions_info


def extract_qa_info(result: Dict[str, Any], qa_version: str = "v3.5") -> Dict[str, Any]:
    """ QA """
    qa_data = result.get("qa", {})
    
    #  QA 
    qa_key = qa_version.replace(".", "_")
    qa_result = qa_data.get(qa_key) or qa_data.get("v3_5") or qa_data.get("v3_4") or {}
    
    structural_risk = qa_result.get("structural_risk_components", {}) or {}
    guideline = qa_result.get("guideline", {}) or {}
    checks = qa_result.get("checks", {}) or {}
    ranking_sanity = checks.get("ranking_sanity_v3_5", {}) or {}
    stability = ranking_sanity.get("stability_analysis", {}) or {}
    
    errors = qa_result.get("errors", []) or []
    warnings = qa_result.get("warnings", []) or []
    
    # 
    warnings_list = []
    for w in warnings:
        if isinstance(w, dict):
            level = w.get("level", "unknown")
            category = w.get("category", "unknown")
            message = w.get("message", "")
            warnings_list.append(f"- **[{category} · {level}]** {message}")
        else:
            warnings_list.append(f"- {w}")
    
    # 
    guideline_flags = []
    for flag in guideline.get("flags", []):
        guideline_flags.append(
            f"- **{flag.get('id', 'UNKNOWN')}** ({flag.get('level', 'unknown')}): "
            f"{flag.get('message', '')} [: {flag.get('value', 0):.2f}]"
        )
    
    return {
        "qa_status": "✅ " if qa_result.get("ok", False) else "❌ ",
        "qa_traffic_light": guideline.get("traffic_light", "unknown"),
        "fr2_risk": structural_risk.get("fr2_hydrophilic_patch_risk", 0.0),
        "graft_risk": structural_risk.get("grafting_interface_risk", 0.0),
        "anchor_risk": structural_risk.get("cdr3_anchor_risk", 0.0),
        "total_risk": structural_risk.get("total_risk", 0.0),
        "fr2_level": _classify_risk_level(structural_risk.get("fr2_hydrophilic_patch_risk", 0.0)),
        "graft_level": _classify_risk_level(structural_risk.get("grafting_interface_risk", 0.0)),
        "anchor_level": _classify_risk_level(structural_risk.get("cdr3_anchor_risk", 0.0)),
        "total_level": _classify_risk_level(structural_risk.get("total_risk", 0.0)),
        "ranking_stable": "" if stability.get("is_stable", False) else "",
        "stability_score": stability.get("stability_score", 0.0),
        "swap_risk": stability.get("swap_risk", 0.0),
        "ranking_tier": stability.get("tier", "A"),
        "recommended_output_mode": stability.get("recommended_output_mode", "single_lead"),
        "ranking_tier_reason": stability.get("tier_reason", ""),
        "consistency_status": "" if not stability.get("consistency_issues") else "",
        "error_count": len(errors),
        "qa_errors_list": "\n".join([f"- {e}" for e in errors]) if errors else "",
        "warning_count": len(warnings),
        "qa_warnings_list": "\n".join(warnings_list) if warnings_list else "",
        "qa_guideline_flags": "\n".join(guideline_flags) if guideline_flags else "",
    }


def _classify_risk_level(value: float) -> str:
    """"""
    if value <= 0.2:
        return ""
    elif value < 0.6:
        return ""
    else:
        return ""


def extract_panel_info(result: Dict[str, Any], panel: str) -> Dict[str, Any]:
    """（A/B/C）"""
    #  result 
    #  result  panels 
    panels = result.get("panels", {}) or {}
    panel_data = panels.get(panel.upper(), {}) or {}
    
    best_match = result.get("best_match", {}) or {}
    humanized = best_match.get("humanized_sequence", "")
    mutations = result.get("mutations", {}).get("list", []) or []
    
    # 
    mutations_table = []
    for m in mutations[:10]:  # 10
        mutations_table.append(
            f"| {m.get('region', 'N/A')} | {m.get('position', 'N/A')} | "
            f"{m.get('from', 'N/A')} → {m.get('to', 'N/A')} |"
        )
    
    mutations_table_str = "\n".join(mutations_table) if mutations_table else ""
    
    qa_data = result.get("qa", {})
    qa_ok = qa_data.get("ok", False)
    
    return {
        f"panel_{panel.lower()}_name": f" {panel.upper()}",
        f"panel_{panel.lower()}_strategy": panel_data.get("strategy", ""),
        f"panel_{panel.lower()}_fr_identity": panel_data.get("fr_identity", 0.0),
        f"panel_{panel.lower()}_developability_score": panel_data.get("developability", {}).get("score", 0.0),
        f"panel_{panel.lower()}_immunogenicity_risk": panel_data.get("immunogenicity", {}).get("risk", "unknown"),
        f"panel_{panel.lower()}_sequence": humanized,
        f"panel_{panel.lower()}_mutations_table": mutations_table_str,
        f"panel_{panel.lower()}_qa_status": "✅ " if qa_ok else "❌ ",
    }


def extract_template_info(result: Dict[str, Any]) -> Dict[str, Any]:
    """"""
    best_match = result.get("best_match", {}) or {}
    template = best_match.get("template", {}) or {}
    candidates = result.get("candidates", []) or []
    
    best_template_id = template.get("id", "UNKNOWN")
    best_fr_identity = best_match.get("scores", {}).get("fr_identity", 0.0)
    best_combined = best_match.get("scores", {}).get("combined_score", 0.0)
    best_structural_risk = best_match.get("scores", {}).get("structural_risk", 0.0)
    
    # 
    comparison_rows = []
    for i, cand in enumerate(candidates[:5], 1):  # 5
        cand_id = cand.get("template_id", f"Candidate_{i}")
        cand_fr = cand.get("scores", {}).get("fr_identity", 0.0)
        cand_combined = cand.get("scores", {}).get("combined_score", 0.0)
        comparison_rows.append(
            f"| {i} | {cand_id} | {cand_fr:.2%} | {cand_combined:.3f} |"
        )
    
    comparison_table = "\n".join(comparison_rows) if comparison_rows else ""
    
    return {
        "best_template_id": best_template_id,
        "best_fr_identity": best_fr_identity,
        "best_combined_score": best_combined,
        "best_structural_risk": best_structural_risk,
        "best_template_sequence": template.get("sequence", "N/A"),
        "template_selection_reasoning": (
            f" FR identity ({best_fr_identity:.2%})  ({best_combined:.3f}) "
            f"。 {best_structural_risk:.3f}。"
        ),
        "template_comparison_table": comparison_table,
    }


def extract_cmc_info(result: Dict[str, Any]) -> Dict[str, Any]:
    """ CMC """
    cmc = result.get("cmc", {}) or {}
    original_cmc = cmc.get("original", {}) or {}
    humanized_cmc = cmc.get("humanized", {}) or {}
    
    return {
        "original_cmc_risk": original_cmc.get("risk_level", "unknown"),
        "humanized_cmc_risk": humanized_cmc.get("risk_level", "unknown"),
        "cmc_risk_delta": "" if humanized_cmc.get("score", 0) < original_cmc.get("score", 0) else "",
        "cmc_risk_details_table": " CMC ",
        "cmc_high_risk_positions": " CMC ",
    }


def extract_immunogenicity_info(result: Dict[str, Any]) -> Dict[str, Any]:
    """"""
    immuno = result.get("immunogenicity", {}) or {}
    original_immuno = immuno.get("original", {}) or {}
    humanized_immuno = immuno.get("humanized", {}) or {}
    
    return {
        "original_immuno_risk": original_immuno.get("risk_level", "unknown"),
        "humanized_immuno_risk": humanized_immuno.get("risk_level", "unknown"),
        "immuno_risk_delta": "" if humanized_immuno.get("score", 0) < original_immuno.get("score", 0) else "",
        "high_risk_epitopes_table": "",
    }


def extract_developability_info(result: Dict[str, Any]) -> Dict[str, Any]:
    """"""
    dev = result.get("developability", {}) or {}
    original_dev = dev.get("original", {}) or {}
    humanized_dev = dev.get("humanized", {}) or {}
    
    return {
        "original_dev_score": original_dev.get("score", 0.0),
        "humanized_dev_score": humanized_dev.get("score", 0.0),
        "dev_delta": humanized_dev.get("score", 0.0) - original_dev.get("score", 0.0),
        "orig_agg": original_dev.get("aggregation", "N/A"),
        "hum_agg": humanized_dev.get("aggregation", "N/A"),
        "delta_agg": "N/A",
        "orig_sol": original_dev.get("solubility", "N/A"),
        "hum_sol": humanized_dev.get("solubility", "N/A"),
        "delta_sol": "N/A",
        "orig_stab": original_dev.get("stability", "N/A"),
        "hum_stab": humanized_dev.get("stability", "N/A"),
        "delta_stab": "N/A",
    }


def build_repair_comparison_rows(comparison: Dict[str, Any]) -> str:
    """
    
    
    Args:
        comparison: comparison，baselinecandidates
    
    Returns:
        Markdown
    """
    if not comparison:
        return "。"
    
    candidates = comparison.get("candidates", []) or []
    if not candidates:
        return "。"
    
    rows = []
    for c in candidates:
        def fmt_pair(b, x, reverse_better=False):
            """→，reverse_better=True"""
            if b is None or x is None:
                return "NA"
            if reverse_better:
                arrow = "↓" if x < b else ("↑" if x > b else "→")
            else:
                arrow = "↑" if x > b else ("↓" if x < b else "→")
            
            # 
            if isinstance(b, float) and isinstance(x, float):
                return f"{b:.2f} → {x:.2f} {arrow}"
            else:
                return f"{b} → {x} {arrow}"
        
        tr = fmt_pair(c.get("baseline_total_risk"), c.get("candidate_total_risk"), reverse_better=True)
        ar = fmt_pair(c.get("baseline_agg_risk"), c.get("candidate_agg_risk"), reverse_better=True)
        cmc = fmt_pair(
            c.get("baseline_cmc_hotspots"), c.get("candidate_cmc_hotspots"), reverse_better=True
        )
        imm = fmt_pair(
            c.get("baseline_immunogenicity"), c.get("candidate_immunogenicity"), reverse_better=True
        )
        kd = fmt_pair(c.get("baseline_kd"), c.get("candidate_kd"), reverse_better=False)  # KD，
        
        scaffold_id = c.get("scaffold_id", "UNKNOWN")
        rows.append(
            f"| {scaffold_id} | {tr} | {ar} | {cmc} | {imm} | {kd} |"
        )
    
    return "\n".join(rows)


def _build_repair_comparison_section(result: Dict[str, Any]) -> str:
    """
    
    
    Args:
        result: 
    
    Returns:
        Markdown
    """
    comparison = result.get("comparison", {}) or {}
    baseline = comparison.get("baseline", {}) or {}
    
    if not baseline and not comparison.get("candidates"):
        return "。"
    
    lines = []
    lines.append("### 12.1 （clone）")
    lines.append("")
    
    # 
    baseline_metrics = [
        ("Total structural risk", baseline.get("total_risk")),
        ("FR2 risk", baseline.get("fr2_risk")),
        ("CDR3 anchor risk", baseline.get("anchor_risk")),
        ("Aggregation risk", baseline.get("agg_risk")),
        ("CMC hotspots ", baseline.get("cmc_hotspot_count")),
        ("Immunogenicity score", baseline.get("immunogenicity_score")),
        (" (KD)", baseline.get("predicted_affinity")),
    ]
    
    for label, value in baseline_metrics:
        if value is not None:
            if isinstance(value, float):
                lines.append(f"- **{label}**：{value:.3f}")
            else:
                lines.append(f"- **{label}**：{value}")
        else:
            lines.append(f"- **{label}**：N/A")
    
    lines.append("")
    lines.append("### 12.2  vs ")
    lines.append("")
    lines.append("| Scaffold | Total risk (base → cand) | Aggregation (base → cand) | CMC hotspots (base → cand) | Immunogenicity (base → cand) | KD (base → cand) |")
    lines.append("|----------|--------------------------|----------------------------|-----------------------------|-------------------------------|-------------------|")
    
    rows = build_repair_comparison_rows(comparison)
    lines.append(rows)
    lines.append("")
    lines.append("> ：clone（↓，↑）。")
    lines.append("")
    lines.append("### 12.3 ")
    lines.append("")
    
    # 
    candidates = comparison.get("candidates", []) or []
    if candidates:
        # （）
        best_candidate = None
        best_score = float('inf')
        
        for c in candidates:
            # ：delta（）
            score = 0
            if c.get("delta_total_risk") is not None:
                score += c.get("delta_total_risk", 0) * 2  # 
            if c.get("delta_agg_risk") is not None:
                score += c.get("delta_agg_risk", 0)
            if c.get("delta_cmc_hotspots") is not None:
                score += c.get("delta_cmc_hotspots", 0) * 0.5
            if c.get("delta_immunogenicity") is not None:
                score += c.get("delta_immunogenicity", 0)
            
            if score < best_score:
                best_score = score
                best_candidate = c
        
        if best_candidate:
            scaffold_id = best_candidate.get("scaffold_id", "UNKNOWN")
            lines.append(f"****：{scaffold_id}")
            lines.append("")
            lines.append("clone：")
            
            improvements = []
            if best_candidate.get("delta_total_risk", 0) < 0:
                improvements.append(f" {abs(best_candidate['delta_total_risk']):.3f}")
            if best_candidate.get("delta_agg_risk", 0) < 0:
                improvements.append(f" {abs(best_candidate['delta_agg_risk']):.3f}")
            if best_candidate.get("delta_cmc_hotspots", 0) < 0:
                improvements.append(f"CMC hotspots  {abs(best_candidate['delta_cmc_hotspots'])}")
            if best_candidate.get("delta_immunogenicity", 0) < 0:
                improvements.append(f" {abs(best_candidate['delta_immunogenicity']):.3f}")
            
            if improvements:
                for imp in improvements:
                    lines.append(f"- {imp}")
            else:
                lines.append("- ，")
        else:
            lines.append("，。")
    else:
        lines.append("，。")
    
    return "\n".join(lines)


def _summarize_features_for_table(features: Dict[str, Any]) -> str:
    """
     affinity.hotspots[i].features ，。
    """
    if not features:
        return "-"
    tags = []
    
    if features.get("is_aromatic"):
        tags.append("")
    if features.get("is_pos_charged"):
        tags.append("")
    if features.get("is_neg_charged"):
        tags.append("")
    if features.get("is_hydrophobic"):
        tags.append("")
    if features.get("is_flexible"):
        tags.append("")
    if features.get("is_rigid"):
        tags.append("")
    
    neighbor_aromatic = features.get("neighbor_aromatic_count", 0)
    if neighbor_aromatic:
        tags.append(f"×{neighbor_aromatic}")
    
    return "，".join(tags) if tags else "-"


def build_affinity_hotspot_table(affinity: Dict[str, Any]) -> str:
    """
     Markdown ， {{affinity.hotspot_table}}。
    """
    rows: List[str] = []
    hotspots = affinity.get("hotspots", []) or []
    
    for h in hotspots[:20]:  # 20
        pos = h.get("position", "")
        aa = h.get("aa", h.get("features", {}).get("aa", ""))
        region = h.get("region", "")
        score = h.get("score", 0.0)
        feat = h.get("features", {}) or {}
        feat_text = _summarize_features_for_table(feat)
        
        rows.append(
            f"| {pos} | {aa} | {region} | {score:.2f} | {feat_text} |"
        )
    
    return "\n".join(rows) if rows else "| - | - | - | - |  |"


def build_affinity_candidate_table(affinity: Dict[str, Any]) -> str:
    """
    ， {{affinity.candidate_table}}。
    """
    rows: List[str] = []
    candidates = affinity.get("candidates", []) or []
    
    for c in candidates[:15]:  # 15
        pos = c.get("position", "")
        from_aa = c.get("from_aa", "")
        to_aa = c.get("to_aa", "")
        net_score = c.get("net_score", 0.0)
        gain = c.get("affinity_gain_score", 0.0)
        penalty = c.get("risk_penalty", 0.0)
        rationale = c.get("rationale", "")
        
        #  rationale 
        if len(rationale) > 50:
            rationale = rationale[:47] + "..."
        
        rows.append(
            f"| {pos} | {from_aa}→{to_aa} | {net_score:.2f} | "
            f"{gain:.2f} | {penalty:.2f} | {rationale} |"
        )
    
    return "\n".join(rows) if rows else "| - | - | - | - | - |  |"


def _build_affinity_variant_table_for_tier(
    affinity: Dict[str, Any],
    tier: str,
) -> str:
    """
     mild / moderate / aggressive 。
    tier: "mild" / "moderate" / "aggressive"
    """
    variants = (affinity.get("variants", {}) or {}).get(tier, []) or []
    if not variants:
        return "_。_"
    
    rows: List[str] = []
    header = (
        "|  |  |  |  |  |\n"
        "|----------|--------|----------|----------------|----------|"
    )
    rows.append(header)
    
    for v in variants:
        name = v.get("name", "")
        muts = v.get("mutations", []) or []
        mut_count = len(muts)
        
        # ：Pos From→To; ...
        mut_items: List[str] = []
        for m in muts:
            pos = m.get("position", "")
            from_aa = m.get("from_aa", "")
            to_aa = m.get("to_aa", "")
            mut_items.append(f"{pos} {from_aa}→{to_aa}")
        mut_text = "; ".join(mut_items)
        
        aff_score = v.get("predicted_affinity_score", 0.0)
        overall = v.get("overall_score", 0.0)
        
        rows.append(
            f"| {name} | {mut_count} | {mut_text} | "
            f"{aff_score:.2f} | {overall:.2f} |"
        )
    
    return "\n".join(rows)


def _build_affinity_optimization_section(result: Dict[str, Any]) -> Dict[str, Any]:
    """
    
    
    Args:
        result: 
    
    Returns:
        
    """
    affinity = result.get("affinity", {}) or {}
    hotspots = affinity.get("hotspots", []) or []
    candidates = affinity.get("candidates", []) or []
    
    return {
        "affinity.hotspot_count": len(hotspots),
        "affinity.hotspot_table": build_affinity_hotspot_table(affinity),
        "affinity.candidate_count": len(candidates),
        "affinity.candidate_table": build_affinity_candidate_table(affinity),
        "affinity.variant_mild_table": _build_affinity_variant_table_for_tier(affinity, "mild"),
        "affinity.variant_moderate_table": _build_affinity_variant_table_for_tier(affinity, "moderate"),
        "affinity.variant_aggressive_table": _build_affinity_variant_table_for_tier(affinity, "aggressive"),
        "affinity.narrative": affinity.get("narrative", "。"),
    }


def build_sections_from_result(result: Dict[str, Any]) -> Dict[str, Any]:
    """
     result 
    
    Args:
        result: 
    
    Returns:
        
    """
    import sys
    from pathlib import Path
    project_root = Path(__file__).parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))
    
    from core.vhh_humanization_panels import build_triple_panel_markdown
    from core.vhh_process_logger import format_process_log_block
    
    # 
    project_id = result.get("project_id", "unknown_project")
    target = result.get("target", result.get("input", {}).get("target", "Unknown"))
    analysis_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    engine_version = result.get("engine_version", "1.0")
    qa_version = result.get("qa", {}).get("active_version", "v3.5")
    
    # 
    input_data = result.get("input", {}) or {}
    original_seq = input_data.get("sequence", "")
    seq_length = len(original_seq)
    
    # （）
    cys_count = original_seq.count("C")
    # pI  GRAVY ，
    pi = "N/A"  #  Bio.SeqUtils.IsoelectricPoint
    gravy = "N/A"  #  Bio.SeqUtils.ProtParam
    
    # IMGT 
    seq_analysis = result.get("sequence_analysis", {}) or {}
    orig_regions = seq_analysis.get("original_regions", {}) or {}
    imgt_rows = []
    for region in ["FR1", "CDR1", "FR2", "CDR2", "FR3", "CDR3", "FR4"]:
        seq = orig_regions.get(region, "")
        if seq:
            imgt_rows.append(f"| {region} | {seq} | {len(seq)} |")
    imgt_table_rows = "\n".join(imgt_rows) if imgt_rows else ""
    
    # 
    qa_info = extract_qa_info(result, qa_version)
    developability_info = extract_developability_info(result)
    cmc_info = extract_cmc_info(result)
    immuno_info = extract_immunogenicity_info(result)
    
    # （，）
    background_section = f" {target}  VHH 。"
    germline_section = "。"
    hallmark_section = " VHH 。"
    developability_section = f"：{developability_info.get('original_dev_score', 'N/A')}，：{developability_info.get('humanized_dev_score', 'N/A')}。"
    immunogenicity_section = f"：{immuno_info.get('original_immuno_risk', 'N/A')}，：{immuno_info.get('humanized_immuno_risk', 'N/A')}。"
    affinity_section = "。"
    
    # QA 
    qa_summary_lines = [
        f"**QA **：{qa_info.get('qa_status', 'N/A')}",
        f"****：{qa_info.get('qa_traffic_light', 'unknown')}",
        f"****：{qa_info.get('total_risk', 0.0):.3f} ({qa_info.get('total_level', 'N/A')})",
        f"****：{qa_info.get('ranking_stable', 'N/A')}",
        "",
        "****：",
        qa_info.get("qa_errors_list", ""),
        "",
        "****：",
        qa_info.get("qa_warnings_list", ""),
    ]
    qa_summary_section = "\n".join(qa_summary_lines)
    
    # 
    triple_panel_md = build_triple_panel_markdown(result)
    
    # 
    process_log = result.get("process_log", []) or []
    process_log_block = format_process_log_block(process_log)
    
    # 
    final_recommendation = result.get("final_recommendation") or (
        "，。"
    )
    
    # 
    repair_comparison_section = _build_repair_comparison_section(result)
    
    # 
    affinity_section_data = _build_affinity_optimization_section(result)
    
    # 
    mapping = {
        "project_id": project_id,
        "target": target,
        "analysis_date": analysis_date,
        "engine_version": engine_version,
        "qa_version": qa_version,
        "background_section": background_section,
        "seq_length": seq_length,
        "pi": pi,
        "gravy": gravy,
        "cys_count": cys_count,
        "input_sequence": original_seq,
        "imgt_table_rows": imgt_table_rows,
        "germline_section": germline_section,
        "hallmark_section": hallmark_section,
        "triple_panel_markdown": triple_panel_md,
        "developability_section": developability_section,
        "immunogenicity_section": immunogenicity_section,
        "affinity_section": affinity_section,
        "qa_summary_section": qa_summary_section,
        "process_log_block": process_log_block,
        "final_recommendation": final_recommendation,
        "repair_comparison_section": repair_comparison_section,
    }
    
    # （）
    mapping.update(affinity_section_data)
    
    return mapping


def render_template(template: str, context: Dict[str, Any]) -> str:
    """，（ {{key}} ）"""
    content = template
    
    #  {{key}} 
    for key, value in context.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in content:
            content = content.replace(placeholder, str(value))
    
    return content


def save_markdown(content: str, output_path: Path):
    """ Markdown """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")
    print(f"✅ Markdown report saved: {output_path}")


def save_docx(content: str, output_path: Path):
    """ DOCX （ python-docx）"""
    if not HAS_DOCX:
        print("⚠️  python-docx not installed, skipping .docx export")
        return
    
    #  Markdown  DOCX
    doc = Document()
    
    # 
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        
        if not line:
            # 
            doc.add_paragraph()
        elif line.startswith("# "):
            # 
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            # 
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            # 
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("|"):
            # （）
            p = doc.add_paragraph(line)
        elif line.startswith("- ") or line.startswith("* "):
            # 
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("```"):
            # （）
            continue
        else:
            # 
            p = doc.add_paragraph(line)
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ DOCX report saved: {output_path}")


def generate_report(
    result_json_path: Path,
    template_path: Path,
    output_dir: Path,
    project_id: Optional[str] = None,
) -> Dict[str, Path]:
    """
    
    
    Returns:
        
    """
    #  result.json
    with open(result_json_path, "r", encoding="utf-8") as f:
        result = json.load(f)
    
    #  project_id
    if not project_id:
        project_id = result.get("project_id", "unknown_project")
    
    # 
    template = load_template(template_path)
    
    # 
    context = {
        "project_id": project_id,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "qa_version": "3.5",
        "suite_version": "1.0",
        "overall_status": "✅ " if result.get("status") == "OK" else "❌ ",
        "input_sequence_name": result.get("input", {}).get("sequence_name", "Unknown"),
        "original_sequence": result.get("input", {}).get("sequence", ""),
        "original_length": len(result.get("input", {}).get("sequence", "")),
        "species": result.get("input", {}).get("species", "alpaca"),
    }
    
    # 
    context = build_sections_from_result(result)
    
    # （ {{key}} ）
    rendered = render_template(template, context)
    
    # 
    output_project_dir = output_dir / project_id
    output_project_dir.mkdir(parents=True, exist_ok=True)
    (output_project_dir / "figures").mkdir(exist_ok=True)
    
    # 
    report_md_path = output_project_dir / "report_v1.md"
    report_docx_path = output_project_dir / "report_v1.docx"
    
    save_markdown(rendered, report_md_path)
    save_docx(rendered, report_docx_path)
    
    return {
        "markdown": report_md_path,
        "docx": report_docx_path,
        "output_dir": output_project_dir,
    }


def main():
    """"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate VHH humanization report v1.0")
    parser.add_argument("--input", "-i", type=Path, required=True, help="Path to result.json")
    parser.add_argument("--template", "-t", type=Path, default=None, help="Path to template (default: reports/templates/vhh_full_report_template.md)")
    parser.add_argument("--output", "-o", type=Path, default=None, help="Output directory (default: reports/output)")
    parser.add_argument("--project-id", "-p", type=str, default=None, help="Project ID (default: from result.json)")
    
    args = parser.parse_args()
    
    # 
    if args.template is None:
        args.template = Path(__file__).parent.parent / "reports" / "templates" / "vhh_full_report_template.md"
    
    output_dir = args.output if args.output else Path(__file__).parent.parent / "reports" / "output"
    
    # 
    try:
        result = generate_report(
            result_json_path=args.input,
            template_path=args.template,
            output_dir=output_dir,
            project_id=args.project_id,
        )
        
        print("\n" + "="*60)
        print("✅ Report generation completed!")
        print(f"[INFO] Markdown report: {result['markdown']}")
        print(f"[INFO] DOCX report    : {result['docx']}")
        print("="*60)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())

