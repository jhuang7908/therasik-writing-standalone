"""
Submission-ready DOCX formatter — journal peer-review submission standard.

Produces TWO documents matching what journals actually require at submission:

  Document 1: TITLE PAGE (title_page.docx)
    - Full title, running head
    - Authors with affiliation superscripts
    - Full affiliations, ORCID, corresponding author email
    - Keywords, word counts (abstract / main text / total)
    - Figure count, table count
    - CRediT author contributions
    - Funding, competing interests, ethics, data availability

  Document 2: BLIND MANUSCRIPT (manuscript.docx)
    - Running head (right header) + page number (right header)
    - Title ONLY (no authors — blinded for peer review)
    - Abstract (unstructured OR structured: Background/Methods/Results/Conclusions)
    - Keywords
    - Body sections with [Figure X near here] placement markers
    - References (hanging-indent, numbered)
    - Figure Legends (after references, on new page)
    - Tables (after figure legends, or flagged for separate upload)

Format spec (matches ICMJE / most major journal requirements):
  - Times New Roman 12pt (or Arial 12pt)
  - Double-spaced throughout
  - 25.4 mm (1-inch) margins
  - Continuous line numbers
  - Page numbers (top-right, same line as running head)
  - Hanging indent for references (0.5 inch)

Public API:
    available() -> bool
    format_title_page(manuscript, options) -> bytes     (DOCX)
    format_blind_manuscript(manuscript, options) -> bytes (DOCX)
    format_manuscript(manuscript, **options) -> bytes   (legacy: blind manuscript only)
    preflight(manuscript) -> dict                       (submission checklist)
"""
from __future__ import annotations

import io
import re
from typing import Any

try:
    from docx import Document
    from docx.shared import Cm, Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ─── Constants ────────────────────────────────────────────────────────────────

RUNNING_HEAD_MAX = 50      # characters (most journals)
ABSTRACT_WORD_LIMIT = 300  # warn above this (most journals: 150–300)

_CREDIT_ROLES = [
    "Conceptualization", "Data curation", "Formal analysis",
    "Funding acquisition", "Investigation", "Methodology",
    "Project administration", "Resources", "Software",
    "Supervision", "Validation", "Visualization",
    "Writing – original draft", "Writing – review & editing",
]

# ─── Availability ─────────────────────────────────────────────────────────────

def available() -> bool:
    return _DOCX_AVAILABLE


# ─── Word counting ────────────────────────────────────────────────────────────

def _word_count(text: str) -> int:
    return len(re.findall(r"\w+", text))


def _count_manuscript_words(manuscript: dict[str, Any]) -> dict[str, int]:
    abstract = manuscript.get("abstract_text") or ""
    sections = manuscript.get("sections") or []
    body_text = " ".join((s.get("text") or "") for s in sections)
    return {
        "abstract": _word_count(abstract),
        "body": _word_count(body_text),
        "total": _word_count(abstract) + _word_count(body_text),
    }


# ─── Low-level DOCX helpers ───────────────────────────────────────────────────

def _new_doc(font_name: str = "Times New Roman", font_size_pt: float = 12.0) -> Any:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    return doc


def _set_margins(section, cm: float = 2.54) -> None:
    m = Cm(cm)
    section.top_margin = m
    section.bottom_margin = m
    section.left_margin = m
    section.right_margin = m


def _add_line_numbers(doc: Any) -> None:
    try:
        sect = doc.sections[0]
        lnNumEl = OxmlElement("w:lnNumType")
        lnNumEl.set(qn("w:countBy"), "1")
        lnNumEl.set(qn("w:restart"), "newPage")
        lnNumEl.set(qn("w:distance"), "720")
        sect._sectPr.append(lnNumEl)
    except Exception:
        pass


def _add_running_head_with_page(doc: Any, running_head: str, font_name: str) -> None:
    """
    Standard journal header: RUNNING HEAD (left-aligned caps) + page number (right).
    Appears on every page including first.
    """
    try:
        section = doc.sections[0]
        section.different_first_page_header_footer = False
        header = section.header
        para = header.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Running head text (left portion via tab stop)
        rh = running_head[:RUNNING_HEAD_MAX].upper()
        r1 = para.add_run(rh + "    ")
        r1.font.name = font_name
        r1.font.size = Pt(10)
        # Page number field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        r2 = para.add_run()
        r2.font.name = font_name
        r2.font.size = Pt(10)
        r2._r.append(fldChar1)
        r2._r.append(instrText)
        r2._r.append(fldChar2)
    except Exception:
        pass


def _para(doc: Any, text: str = "", *,
          font_name: str = "Times New Roman",
          size: float = 12,
          bold: bool = False,
          italic: bool = False,
          align: Any = None,
          space_before: float = 0,
          space_after: float = 0,
          color: RGBColor | None = None,
          hanging: float = 0,
          left_indent: float = 0) -> Any:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if hanging:
        pf.first_line_indent = Pt(-hanging)
        pf.left_indent = Pt(hanging)
    elif left_indent:
        pf.left_indent = Pt(left_indent)
    if text:
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def _heading(doc: Any, text: str, font_name: str = "Times New Roman",
             level: int = 1) -> Any:
    """Level 1: ALL CAPS bold. Level 2: Title Case bold italic (subsection)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    display = text.upper() if level == 1 else text
    r = p.add_run(display)
    r.font.name = font_name
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.italic = (level == 2)
    return p


def _citation_para(doc: Any, text: str, font_name: str) -> Any:
    """Body paragraph with inline [1] or (Author, Year) superscript citation markers."""
    pattern = re.compile(
        r"(\[\d+(?:[,;\s]*\d+)*\]"          # [1] [1,2] [1; 2]
        r"|\([A-Z][a-zA-Z\-]+(?:\s+et\s+al\.)?,\s*\d{4}[a-z]?\))"  # (Smith et al., 2024)
    )
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for part in pattern.split(text):
        if part and pattern.fullmatch(part):
            r = p.add_run(part)
            r.font.name = font_name
            r.font.size = Pt(10)
            r.font.superscript = True
        elif part:
            r = p.add_run(part)
            r.font.name = font_name
            r.font.size = Pt(12)
    return p


def _figure_placeholder(doc: Any, n: int, font_name: str) -> Any:
    """Standard journal in-text figure placement marker."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[Insert Figure {n} about here]")
    r.font.name = font_name
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def _table_placeholder(doc: Any, n: int, caption: str, font_name: str) -> Any:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[Insert Table {n} about here: {caption}]")
    r.font.name = font_name
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def _page_break(doc: Any) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ─── Structured abstract builder ─────────────────────────────────────────────

def _render_abstract(
    doc: Any,
    abstract: str,
    structured: bool,
    font_name: str,
) -> None:
    """Render abstract section — plain or structured (Background/Methods/Results/Conclusions)."""
    _heading(doc, "Abstract", font_name)
    if not abstract:
        _para(doc, "[Abstract not provided]", font_name=font_name,
              italic=True, color=RGBColor(0x99, 0x99, 0x99))
        return

    if not structured:
        _citation_para(doc, abstract.strip(), font_name)
        return

    # Auto-detect or split structured abstract
    STRUCTURED_HEADINGS = ["Background", "Objective", "Methods", "Results",
                            "Conclusions", "Conclusion", "Significance",
                            "Importance", "Purpose", "Design", "Setting",
                            "Participants", "Interventions", "Main Outcomes"]
    heading_pat = re.compile(
        r"^(" + "|".join(STRUCTURED_HEADINGS) + r")\s*[:\.]?\s*",
        re.IGNORECASE | re.MULTILINE,
    )
    parts = heading_pat.split(abstract.strip())
    if len(parts) <= 1:
        # No subheadings found — render plain
        _citation_para(doc, abstract.strip(), font_name)
        return

    # Render each sub-section: bold label inline with text
    i = 1
    while i < len(parts):
        label = parts[i].strip().rstrip(":")
        body = parts[i + 1].strip() if (i + 1) < len(parts) else ""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r_label = p.add_run(label + ": ")
        r_label.font.name = font_name
        r_label.font.size = Pt(12)
        r_label.font.bold = True
        r_body = p.add_run(body)
        r_body.font.name = font_name
        r_body.font.size = Pt(12)
        i += 2


# ─── Document 1: Title Page ───────────────────────────────────────────────────

def format_title_page(
    manuscript: dict[str, Any],
    *,
    font_name: str = "Times New Roman",
    font_size_pt: float = 12.0,
) -> bytes:
    """
    Produce the TITLE PAGE DOCX (submitted separately from the blind manuscript).
    Contains all author-identifying information that is stripped from the blind manuscript.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed.")

    doc = _new_doc(font_name, font_size_pt)
    _set_margins(doc.sections[0])

    counts = _count_manuscript_words(manuscript)
    title = manuscript.get("title") or "Untitled Manuscript"
    article_type = (manuscript.get("article_type") or "Research Article").title()
    running_head = manuscript.get("running_head") or title[:RUNNING_HEAD_MAX]

    # ── Title ────────────────────────────────────────────────────────────────
    _para(doc, title, font_name=font_name, size=14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    _para(doc, f"Article type: {article_type}",
          font_name=font_name, size=10, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55))

    _para(doc, f"Running head: {running_head.upper()[:RUNNING_HEAD_MAX]}",
          font_name=font_name, size=10, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55),
          space_after=12)

    # ── Authors ───────────────────────────────────────────────────────────────
    authors_raw = manuscript.get("authors") or ""
    if authors_raw:
        _para(doc, str(authors_raw), font_name=font_name, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

    # ── Affiliations ──────────────────────────────────────────────────────────
    affiliations = manuscript.get("affiliations") or ""
    if affiliations:
        _para(doc, str(affiliations), font_name=font_name, size=10, italic=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

    # ── Corresponding author ──────────────────────────────────────────────────
    corr = manuscript.get("corresponding_author") or {}
    corr_text = (
        f"Corresponding author: {corr.get('name','[FILL: name]')}, "
        f"{corr.get('institution','[FILL: institution]')}, "
        f"Email: {corr.get('email','[FILL: email]')}"
    )
    if corr.get("orcid"):
        corr_text += f", ORCID: {corr['orcid']}"
    _para(doc, corr_text, font_name=font_name, size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ── Word counts ────────────────────────────────────────────────────────────
    figs = len(manuscript.get("figure_legends") or [])
    tables = len(manuscript.get("tables") or [])
    _heading(doc, "Manuscript Statistics", font_name)
    stats_lines = [
        f"Abstract word count: {counts['abstract']} words",
        f"Main text word count: {counts['body']} words (excluding abstract, references, legends)",
        f"Total word count: {counts['total']} words",
        f"Number of figures: {figs}",
        f"Number of tables: {tables}",
        f"Number of references: {len(manuscript.get('reference_list') or [])}",
    ]
    for line in stats_lines:
        _para(doc, line, font_name=font_name)

    # ── Keywords ─────────────────────────────────────────────────────────────
    keywords = manuscript.get("keywords") or ""
    if keywords:
        _heading(doc, "Keywords", font_name)
        _para(doc, str(keywords), font_name=font_name)

    # ── CRediT Author Contributions ───────────────────────────────────────────
    credit = manuscript.get("author_contributions") or {}
    if credit:
        _heading(doc, "Author Contributions", font_name)
        if isinstance(credit, str):
            _para(doc, credit, font_name=font_name)
        elif isinstance(credit, dict):
            for author, roles in credit.items():
                if isinstance(roles, list):
                    roles = "; ".join(roles)
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                r1 = p.add_run(f"{author}: ")
                r1.font.name = font_name
                r1.font.size = Pt(12)
                r1.font.bold = True
                r2 = p.add_run(str(roles))
                r2.font.name = font_name
                r2.font.size = Pt(12)
        _para(doc, f"CRediT taxonomy roles available at: https://credit.niso.org/",
              font_name=font_name, size=9, italic=True,
              color=RGBColor(0x88, 0x88, 0x88))

    # ── Declarations ─────────────────────────────────────────────────────────
    declarations = manuscript.get("declarations") or {}
    _heading(doc, "Declarations", font_name)
    decl_defaults = {
        "funding": "[FILL: funding sources and grant numbers, or 'This research received no external funding']",
        "competing_interests": "[FILL: competing interests, or 'The authors declare no competing interests']",
        "ethics_approval": "[FILL: IRB/ethics committee approval number, or 'Not applicable']",
        "consent_to_participate": "[FILL: consent statement, or 'Not applicable']",
        "consent_for_publication": "[FILL: consent statement, or 'Not applicable']",
        "data_availability": "[FILL: data repository DOI or policy, or 'Data available on reasonable request']",
        "acknowledgements": "[FILL: acknowledgements, or none]",
    }
    for key, default in decl_defaults.items():
        val = declarations.get(key) or default
        label = key.replace("_", " ").title()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        r1 = p.add_run(f"{label}: ")
        r1.font.name = font_name
        r1.font.size = Pt(12)
        r1.font.bold = True
        is_fill = str(val).startswith("[FILL")
        r2 = p.add_run(str(val))
        r2.font.name = font_name
        r2.font.size = Pt(12)
        if is_fill:
            r2.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Document 2: Blind Manuscript ────────────────────────────────────────────

def format_blind_manuscript(
    manuscript: dict[str, Any],
    *,
    font_name: str = "Times New Roman",
    font_size_pt: float = 12.0,
    structured_abstract: bool = False,
    line_numbers: bool = True,
    insert_figure_placeholders: bool = False,
) -> bytes:
    """
    Produce the BLIND MANUSCRIPT DOCX (the version uploaded for peer review).
    - No author names (blinded)
    - Running head + page number in header
    - Continuous line numbers
    - Figure placement markers in text
    - Figure Legends on a new page after References
    - Tables on a new page after Figure Legends
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed.")

    doc = _new_doc(font_name, font_size_pt)
    section = doc.sections[0]
    _set_margins(section)
    if line_numbers:
        _add_line_numbers(doc)

    title = manuscript.get("title") or "Untitled Manuscript"
    running_head = manuscript.get("running_head") or title[:RUNNING_HEAD_MAX]
    counts = _count_manuscript_words(manuscript)

    # Running head + page number in header
    _add_running_head_with_page(doc, running_head, font_name)

    # ── Title (no authors — blinded) ──────────────────────────────────────────
    _para(doc, title, font_name=font_name, size=14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Word count display (required by many journals)
    wc_text = (
        f"Word count: {counts['body']} (main text)  |  "
        f"Abstract: {counts['abstract']}  |  "
        f"Figures: {len(manuscript.get('figure_legends') or [])}  |  "
        f"Tables: {len(manuscript.get('tables') or [])}"
    )
    _para(doc, wc_text, font_name=font_name, size=9,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x55, 0x55, 0x55), space_after=12)

    # ── Abstract ─────────────────────────────────────────────────────────────
    abstract = manuscript.get("abstract_text") or ""
    _render_abstract(doc, abstract, structured=structured_abstract, font_name=font_name)

    # ── Keywords ─────────────────────────────────────────────────────────────
    keywords = manuscript.get("keywords") or ""
    if keywords:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        r_kw = p.add_run("Keywords: ")
        r_kw.font.name = font_name
        r_kw.font.size = Pt(12)
        r_kw.font.bold = True
        r_kws = p.add_run(str(keywords))
        r_kws.font.name = font_name
        r_kws.font.size = Pt(12)

    # ── Body sections ─────────────────────────────────────────────────────────
    sections = manuscript.get("sections") or []
    figure_counter = [0]  # mutable so inline closure can increment

    for sec in sections:
        heading_text = (sec.get("title") or sec.get("key") or "Section").strip()
        text = (sec.get("text") or "").strip()
        if not text:
            continue

        _heading(doc, heading_text, font_name)

        # Split text into paragraphs at blank lines
        raw_paras = re.split(r"\n\s*\n+", text)
        for para_text in raw_paras:
            para_text = para_text.strip()
            if not para_text:
                continue
            _citation_para(doc, para_text, font_name)

    # ── References ────────────────────────────────────────────────────────────
    reference_list = manuscript.get("reference_list") or []
    if reference_list:
        _heading(doc, "References", font_name)
        for i, ref in enumerate(reference_list, start=1):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            pf.space_before = Pt(0)
            pf.space_after = Pt(3)
            pf.first_line_indent = Pt(-28)
            pf.left_indent = Pt(28)
            r_n = p.add_run(f"{i}.\u2002")  # en-space after number
            r_n.font.name = font_name
            r_n.font.size = Pt(12)
            r_n.font.bold = False
            r_t = p.add_run(str(ref))
            r_t.font.name = font_name
            r_t.font.size = Pt(12)

    # ── Figure Legends (new page, after references) ───────────────────────────
    # Note: Figure images are prepared separately by the author (e.g., PPT/Illustrator)
    # and uploaded individually in the journal submission system.
    # Only the text legends appear here, as required by submission standards.
    figure_legends = manuscript.get("figure_legends") or []
    if figure_legends:
        _page_break(doc)
        _heading(doc, "Figure Legends", font_name)
        for leg in figure_legends:
            if not isinstance(leg, dict):
                continue
            num = leg.get("figure_number", "?")
            fig_title = leg.get("title") or ""
            full = leg.get("rendered_full") or leg.get("legend_text") or ""
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            pf.space_before = Pt(6)
            pf.space_after = Pt(0)
            pf.first_line_indent = Pt(-28)
            pf.left_indent = Pt(28)
            r_lbl = p.add_run(f"Figure {num}. ")
            r_lbl.font.name = font_name
            r_lbl.font.size = Pt(12)
            r_lbl.font.bold = True
            if fig_title:
                r_tit = p.add_run(f"{fig_title}. ")
                r_tit.font.name = font_name
                r_tit.font.size = Pt(12)
                r_tit.font.bold = True
            if full:
                r_full = p.add_run(full)
                r_full.font.name = font_name
                r_full.font.size = Pt(12)

    # ── Tables (new page, after figure legends) ──────────────────────────────
    tables = manuscript.get("tables") or []
    if tables:
        _page_break(doc)
        _heading(doc, "Tables", font_name)
        for i, tbl in enumerate(tables, start=1):
            _para(doc,
                  f"Table {i}. {tbl.get('title','')}",
                  font_name=font_name, bold=True, space_before=12)
            if tbl.get("caption"):
                _para(doc, tbl["caption"], font_name=font_name, size=10)
            if tbl.get("content"):
                _para(doc, tbl["content"], font_name=font_name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Pre-submission checklist ─────────────────────────────────────────────────

def preflight(manuscript: dict[str, Any]) -> dict[str, Any]:
    """
    Run a pre-submission checklist and return a structured report.
    Does not raise — caller decides how to handle warnings/errors.
    """
    issues: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []

    def error(field: str, msg: str):
        issues.append({"field": field, "message": msg})

    def warn(field: str, msg: str):
        warnings.append({"field": field, "message": msg})

    title = manuscript.get("title") or ""
    if not title.strip():
        error("title", "Title is required.")
    elif len(title) > 300:
        warn("title", f"Title is {len(title)} chars — most journals limit to 150–200 chars.")

    abstract = manuscript.get("abstract_text") or ""
    if not abstract.strip():
        error("abstract_text", "Abstract is required.")
    else:
        wc = _word_count(abstract)
        if wc > ABSTRACT_WORD_LIMIT:
            warn("abstract_text", f"Abstract is {wc} words — most journals limit to 150–300.")
        if wc < 50:
            warn("abstract_text", f"Abstract is only {wc} words — may be too short.")

    keywords = manuscript.get("keywords") or ""
    if not keywords.strip():
        warn("keywords", "Keywords missing — required by most journals.")
    else:
        kw_count = len(re.split(r"[,;]+", str(keywords)))
        if kw_count < 3:
            warn("keywords", f"Only {kw_count} keyword(s) — journals typically require 4–8.")
        elif kw_count > 10:
            warn("keywords", f"{kw_count} keywords — most journals limit to 5–8.")

    sections = manuscript.get("sections") or []
    if not sections:
        error("sections", "No body sections found.")
    else:
        has_intro = any(re.search(r"intro", (s.get("title") or ""), re.I) for s in sections)
        has_methods = any(re.search(r"method|material", (s.get("title") or ""), re.I) for s in sections)
        has_results = any(re.search(r"result|finding", (s.get("title") or ""), re.I) for s in sections)
        has_discussion = any(re.search(r"discussion|conclusion", (s.get("title") or ""), re.I) for s in sections)
        if not has_intro:
            warn("sections", "No Introduction section detected.")
        if not has_methods:
            warn("sections", "No Methods section detected.")
        if not has_results:
            warn("sections", "No Results section detected.")
        if not has_discussion:
            warn("sections", "No Discussion/Conclusions section detected.")

    refs = manuscript.get("reference_list") or []
    if not refs:
        error("reference_list", "No references — required for all journal submissions.")
    elif len(refs) < 5:
        warn("reference_list", f"Only {len(refs)} reference(s) — seems too few.")

    declarations = manuscript.get("declarations") or {}
    for key in ("funding", "competing_interests"):
        if not declarations.get(key):
            warn(f"declarations.{key}", f"'{key}' declaration missing — required by most journals.")

    if not manuscript.get("corresponding_author"):
        warn("corresponding_author", "Corresponding author details missing (name, email, ORCID).")

    running_head = manuscript.get("running_head") or manuscript.get("title") or ""
    if len(running_head) > RUNNING_HEAD_MAX:
        warn("running_head",
             f"Running head is {len(running_head)} chars — exceeds {RUNNING_HEAD_MAX} char limit. "
             "Provide a short running head via the 'running_head' field.")

    counts = _count_manuscript_words(manuscript)

    # Overall pass/fail
    status = "FAIL" if issues else ("WARN" if warnings else "PASS")
    return {
        "status": status,
        "errors": issues,
        "warnings": warnings,
        "word_counts": counts,
        "figure_count": len(manuscript.get("figure_legends") or []),
        "table_count": len(manuscript.get("tables") or []),
        "reference_count": len(refs),
        "summary": (
            f"{len(issues)} error(s), {len(warnings)} warning(s). "
            f"Main text: {counts['body']} words. "
            f"Abstract: {counts['abstract']} words."
        ),
    }


# ─── Legacy compatibility shim ────────────────────────────────────────────────

def format_manuscript(
    manuscript: dict[str, Any],
    *,
    citation_style: str = "numbered",
    justify: bool = False,
    line_numbers: bool = True,
    font_name: str = "Times New Roman",
    font_size_pt: float = 12.0,
) -> bytes:
    """Legacy entry point — returns blind manuscript bytes (same as format_blind_manuscript)."""
    return format_blind_manuscript(
        manuscript,
        font_name=font_name,
        font_size_pt=font_size_pt,
        line_numbers=line_numbers,
    )


__all__ = [
    "available",
    "format_title_page",
    "format_blind_manuscript",
    "format_manuscript",
    "preflight",
]
