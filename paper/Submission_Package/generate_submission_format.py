"""
Word
：
- 
- Times New Roman 12pt
- 
- 
- 1
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_line_numbering(section):
    """"""
    sectPr = section._sectPr
    lnNumType = OxmlElement('w:lnNumType')
    lnNumType.set(qn('w:countBy'), '1')
    lnNumType.set(qn('w:restart'), 'continuous')
    sectPr.append(lnNumType)

def add_page_number(section):
    """"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_submission_format(doc):
    """"""
    # （1 = 72pt = 914400 EMU）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # 
        add_line_numbering(section)
        add_page_number(section)
    
    # 
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

def parse_markdown_to_docx(md_path, docx_path):
    """MarkdownDOCX"""
    print(f": {md_path} -> {docx_path}")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read
    
    doc = Document
    set_submission_format(doc)
    
    lines = content.split('\n')
    i = 0
    
    # （#）
    title_added = False
    while i < len(lines):
        line = lines[i].rstrip
        if not line:
            i += 1
            continue
        
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip
            p = doc.add_paragraph(title_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_after = Pt(12)
            title_added = True
            i += 1
            break
        i += 1
    
    # 
    while i < len(lines):
        line = lines[i].rstrip
        
        if not line:
            i += 1
            continue
        
        # （^）
        if '^' in line and not line.startswith('#'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(11)
            i += 1
            continue
        
        #  (##)
        if line.startswith('## '):
            heading_text = line[3:].strip
            p = doc.add_heading(heading_text, level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        #  (###)
        if line.startswith('### '):
            heading_text = line[4:].strip
            p = doc.add_heading(heading_text, level=2)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        #  (####)
        if line.startswith('#### '):
            heading_text = line[5:].strip
            p = doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # 
        if line.strip == '---':
            doc.add_paragraph
            i += 1
            continue
        
        # 
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # （[]）
        if re.match(r'^\[\d+\]', line):
            p = doc.add_paragraph(line, style='Normal')
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            if p.runs:
                p.runs[0].font.size = Pt(11)
            i += 1
            continue
        
        # 
        if line and not line.startswith('#'):
            p = doc.add_paragraph
            add_formatted_text(p, line)
            i += 1
            continue
        
        i += 1
    
    doc.save(docx_path)
    print(f"✓ Word: {docx_path}")

def add_formatted_text(paragraph, text):
    """"""
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    from pathlib import Path
    base_dir = Path(__file__).parent
    
    print("=" * 80)
    print("Word")
    print("=" * 80)
    print
    
    md_file = base_dir / "Manuscript_VHH_Humanization.md"
    docx_file = base_dir / "Manuscript_VHH_Humanization.docx"
    
    parse_markdown_to_docx(md_file, docx_file)
    
    print
    print(":")
    print("  ✓ ")
    print("  ✓ Times New Roman 12pt")
    print("  ✓ ")
    print("  ✓ ")
    print("  ✓ 1")
    print("  ✓ 14pt")
    print
    print("=" * 80)
    print("✓ ！")
    print("=" * 80)
