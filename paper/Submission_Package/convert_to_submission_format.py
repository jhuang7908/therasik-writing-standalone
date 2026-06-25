"""
Convert submission package files to final formats for journal submission
"""
import re
import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import zipfile
import os

def parse_markdown_to_docx(md_path, docx_path):
    """Convert Markdown manuscript to DOCX with formatting"""
    print(f"Converting {md_path} to DOCX...")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    # First, handle the title (first non-empty line starting with #)
    title_added = False
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines at the beginning
        if not line:
            i += 1
            continue
        
        # Title (first line starting with # but not ##)
        if not title_added and line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()  # Remove '# '
            p = doc.add_paragraph(title_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            title_added = True
            i += 1
            continue
        
        # Once we have the title, process the rest
        if title_added:
            break
    
    # Now continue with the rest of the document
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Authors (lines with ^)
        if '^' in line and not line.startswith('#'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Heading 1 (##)
        if line.startswith('## '):
            heading_text = line[3:].strip()
            p = doc.add_heading(heading_text, level=1)
            i += 1
            continue
        
        # Heading 2 (###)
        if line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # Heading 3 (####)
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            p = doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        # List item
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Keywords line
        if line.startswith('**Keywords**:'):
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            i += 1
            continue
        
        # Reference (starts with [number])
        if re.match(r'^\[\d+\]', line):
            p = doc.add_paragraph(line, style='Normal')
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.runs[0]
            run.font.size = Pt(10)
            i += 1
            continue
        
        # Regular paragraph
        if line and not line.startswith('#'):
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            i += 1
            continue
        
        i += 1
    
    doc.save(docx_path)
    print(f"✓ DOCX saved to: {docx_path}")

def add_formatted_text(paragraph, text):
    """Add text with bold and italic formatting"""
    # Pattern to match **bold** and *italic*
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code (keep as regular text)
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Regular text
            paragraph.add_run(part)

def create_formatted_excel(csv_path, excel_path, table_title):
    """Convert CSV to formatted Excel with professional styling"""
    print(f"Converting {csv_path} to Excel...")
    
    # Read CSV
    df = pd.read_csv(csv_path)
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    
    # Define styles
    header_font = Font(name='Arial', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    cell_font = Font(name='Arial', size=10)
    cell_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    cell_alignment_center = Alignment(horizontal='center', vertical='center')
    
    border_style = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )
    
    # Add title
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
    title_cell = ws.cell(row=1, column=1)
    title_cell.value = table_title
    title_cell.font = Font(name='Arial', size=14, bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add headers (row 3)
    for col_idx, column_name in enumerate(df.columns, start=1):
        cell = ws.cell(row=3, column=col_idx)
        cell.value = column_name
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border_style
    
    # Add data
    for row_idx, row in enumerate(df.itertuples(index=False), start=4):
        for col_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = value
            cell.font = cell_font
            
            # Center-align numeric columns
            if isinstance(value, (int, float)) or (isinstance(value, str) and value.replace('.', '').replace('-', '').isdigit()):
                cell.alignment = cell_alignment_center
            else:
                cell.alignment = cell_alignment
            
            cell.border = border_style
    
    # Adjust column widths
    for col_idx, column_name in enumerate(df.columns, start=1):
        max_length = max(
            len(str(column_name)),
            df.iloc[:, col_idx-1].astype(str).str.len().max() if len(df) > 0 else 0
        )
        adjusted_width = min(max_length + 3, 50)
        ws.column_dimensions[chr(64 + col_idx)].width = adjusted_width
    
    # Set row heights
    ws.row_dimensions[1].height = 25
    ws.row_dimensions[3].height = 30
    
    wb.save(excel_path)
    print(f"✓ Excel saved to: {excel_path}")

def create_zip_archive(source_dir, zip_path):
    """Create ZIP archive of supplementary materials"""
    print(f"Creating ZIP archive: {zip_path}")
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(source_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, os.path.dirname(source_dir))
                zipf.write(file_path, arcname)
                print(f"  Adding: {arcname}")
    
    print(f"✓ ZIP archive created: {zip_path}")

def main():
    base_dir = Path(__file__).parent
    
    print("=" * 60)
    print("SUBMISSION PACKAGE FORMAT CONVERSION")
    print("=" * 60)
    print()
    
    # 1. Convert Manuscript to DOCX
    print("TASK 1: Converting Manuscript to DOCX")
    print("-" * 60)
    md_file = base_dir / "Manuscript_VHH_Humanization.md"
    docx_file = base_dir / "Manuscript_VHH_Humanization.docx"
    parse_markdown_to_docx(md_file, docx_file)
    print()
    
    # 2. Convert Tables to Excel
    print("TASK 2: Converting Tables to Excel")
    print("-" * 60)
    
    # Table 1
    table1_csv = base_dir / "Tables" / "Table1_Clinical_Landscape.csv"
    table1_excel = base_dir / "Tables" / "Table1_Clinical_Landscape.xlsx"
    create_formatted_excel(
        table1_csv, 
        table1_excel,
        "Table 1. Clinical landscape of 19 VHH therapeutics"
    )
    
    # Table S1
    tableS1_csv = base_dir / "Tables" / "TableS1_Residue_Frequencies.csv"
    tableS1_excel = base_dir / "Tables" / "TableS1_Residue_Frequencies.xlsx"
    create_formatted_excel(
        tableS1_csv,
        tableS1_excel,
        "Table S1. Residue frequencies at key IMGT positions stratified by CDR3 length"
    )
    print()
    
    # 3. Create ZIP of Supplementary Materials
    print("TASK 3: Creating ZIP archive of Supplementary Materials")
    print("-" * 60)
    supp_dir = base_dir.parent / "Supplementary_Materials"
    zip_file = base_dir / "Supplementary_Materials.zip"
    create_zip_archive(supp_dir, zip_file)
    print()
    
    print("=" * 60)
    print("✓ ALL CONVERSION TASKS COMPLETED")
    print("=" * 60)
    print()
    print("Generated files:")
    print(f"  1. {docx_file.name}")
    print(f"  2. {table1_excel.name}")
    print(f"  3. {tableS1_excel.name}")
    print(f"  4. {zip_file.name}")
    print()
    print("Ready for submission!")

if __name__ == "__main__":
    main()
